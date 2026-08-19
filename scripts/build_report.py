"""Build the final four-page StateBench Tech Fellows mini-paper."""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE_RESULTS = ROOT / "results" / "final"
LIVE_RESULTS = ROOT / "results" / "final"
OUT = ROOT / "report" / "StateBench_Report.docx"

# standard_business_brief preset with a named academic-four-page override:
# 10 pt body, 0.75 in side margins, 0.68/0.70 in vertical margins, and
# 1.04 line spacing. The override keeps the mandatory submission under four
# pages while retaining a conventional academic-report type size.
FONT = "Calibri"
PAGE_WIDTH_DXA = 10080  # 7.0 in usable width under the named override
TABLE_INDENT_DXA = 120
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 37, 53)
MUTED = RGBColor(86, 96, 108)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run(run, size: float | None = None, bold: bool | None = None, color=None, italic: bool | None = None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3.2)
    normal.paragraph_format.line_spacing = 1.04

    title = doc.styles["Title"]
    title.font.name = FONT
    title.font.size = Pt(19)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle.font.size = Pt(9.2)
    subtitle.font.italic = False
    subtitle.font.color.rgb = MUTED
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(3)

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1.font.size = Pt(12.2)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h1.paragraph_format.space_before = Pt(7)
    h1.paragraph_format.space_after = Pt(3.5)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2.font.size = Pt(10.4)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BLUE
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2.paragraph_format.space_before = Pt(5)
    h2.paragraph_format.space_after = Pt(2.5)
    h2.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(7.6)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(3)
    caption.paragraph_format.keep_together = True


def paragraph_border_bottom(paragraph, color: str = "2E74B5", size: str = "8") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def page_field(paragraph) -> None:
    set_run(paragraph.add_run("StateBench | "), 8, False, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for item in (begin, instr, separate, text, end):
        run = OxmlElement("w:r")
        run.append(item)
        paragraph._p.append(run)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run(header.add_run("LEBNET TECH FELLOWS 2026 | RESEARCH REPORT"), 7.7, True, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    page_field(footer)


def body(doc: Document, text: str, *, after: float = 3.2, italic: bool = False):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.widow_control = True
    set_run(paragraph.add_run(text), 10, False, INK, italic)
    return paragraph


def labeled_body(doc: Document, label: str, text: str, *, after: float = 3.2):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.widow_control = True
    set_run(paragraph.add_run(label + " "), 10, True, DARK_BLUE)
    set_run(paragraph.add_run(text), 10, False, INK)
    return paragraph


def page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=120, bottom=70, end=120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    old = tcpr.find(qn("w:tcMar"))
    if old is not None:
        tcpr.remove(old)
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = OxmlElement(f"w:{side}")
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")
        margins.append(item)
    tcpr.append(margins)


def set_repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def prevent_row_split(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {PAGE_WIDTH_DXA} DXA, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.insert(0, tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    for tag in ("w:tblInd", "w:tblLayout"):
        old = tblpr.find(qn(tag))
        if old is not None:
            tblpr.remove(old)
    tblind = OxmlElement("w:tblInd")
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    tblpr.append(tblind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)
    for gridcol, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        gridcol.set(qn("w:w"), str(width))
    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths):
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def write_cell(cell, text: str, *, size: float = 8, bold: bool = False, color=INK, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if align is not None:
        paragraph.alignment = align
    set_run(paragraph.add_run(str(text)), size, bold, color)


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)
    return paragraph


def add_abstract(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_together = True
    ppr = paragraph._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), LIGHT_BLUE)
    ppr.append(shade)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "4A5568")
        borders.append(border)
    ppr.append(borders)
    set_run(paragraph.add_run("Abstract. "), 9.5, True, DARK_BLUE)
    set_run(paragraph.add_run(text), 9.5, False, INK)


def add_comparison_table(doc: Document) -> None:
    rows = [
        (
            "Dependency depth",
            "Success falls as depth increases.",
            "Paper-style calibration: easy 2/5, deep 0/5. Qualitatively reproduced.",
        ),
        (
            "Connected distractors",
            "CINs are more harmful than no-extra or disconnected settings.",
            "Paper-style deep and connected both reached 0/5, so the incremental effect was not testable there. Under the compatible protocol, full history fell from 5/5 deep to 3/5 connected.",
        ),
        (
            "State-tracking errors",
            "Using values not yet known is the most common error family.",
            "Sixteen value-not-yet-known events occurred across final traces; explicit errors allowed several runs to recover.",
        ),
        (
            "Known-value restatement",
            "Restating all known values often improves reliability, although gains vary by model.",
            "Restatement improved only from 12/15 to 13/15 and stayed 3/5 on connected tasks. This is a qualified, not large, reproduction of the mitigation effect.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            write_cell(cell, value, size=7.6)
    headers = ["Paper finding", "Original result", "StateBench outcome"]
    for cell, value in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_GRAY)
        write_cell(cell, value, size=7.7, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    table_geometry(table, [1900, 2700, 5480])
    add_caption(doc, "Table 1. Relationship between the published findings and this representative reproduction.")


def add_condition_table(doc: Document) -> None:
    rows = [
        ("Easy", "5", "1", "0", "Basic tool selection and value propagation"),
        ("Deep", "5", "4", "0", "Longer hidden dependency chain"),
        ("Connected distractors", "5", "4", "10 CINs", "Type-compatible irrelevant branches"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            write_cell(cell, value, size=7.7, align=WD_ALIGN_PARAGRAPH.LEFT if index in (0, 4) else WD_ALIGN_PARAGRAPH.CENTER)
    headers = ["Condition", "Core", "Depth", "Distractors", "Purpose"]
    for cell, value in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_GRAY)
        write_cell(cell, value, size=7.7, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    table_geometry(table, [2050, 750, 850, 1450, 4980])
    add_caption(doc, "Table 2. Final paired extension: five frozen graph seeds per condition.")


def add_results_table(doc: Document, summary: pd.DataFrame, include_live: bool) -> None:
    strategies = ["full_history", "paper_restatement", "json_state", "dependency_pruning"]
    if include_live:
        strategies.append("live_dependency_pruning")
    ordered = summary.set_index("strategy").reindex(strategies)
    labels = {
        "full_history": "Full history",
        "paper_restatement": "Paper restatement",
        "json_state": "JSON state",
        "dependency_pruning": "Dependency pruning",
        "live_dependency_pruning": "Live dependency pruning",
    }
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    for strategy, row in ordered.iterrows():
        values = [
            labels[strategy],
            f"{int(row.successes)}/{int(row.trials)}",
            f"{row.avg_prompt_tokens:,.0f}",
            f"{row.avg_completion_tokens:,.0f}",
            f"{row.avg_calls:.2f}",
            f"{row.avg_state_bytes:,.1f}",
            f"{row.success_per_1k_prompt_tokens:.3f}",
        ]
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            write_cell(cell, value, size=7.3, align=WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    headers = ["Strategy", "Success", "Prompt\ntokens", "Completion\ntokens", "Calls", "State\nbytes", "Success /\n1k prompt"]
    for cell, value in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_GRAY)
        write_cell(cell, value, size=7.3, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    table_geometry(table, [2200, 850, 1350, 1400, 850, 1250, 2180])
    add_caption(doc, "Table 3. Aggregate final results across 15 tasks. State bytes are cumulative added-state suffix bytes per trial.")


def add_picture(doc: Document, path: Path, width: float, alt_text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    docpr = inline_shape._inline.docPr
    docpr.set("title", alt_text)
    docpr.set("descr", alt_text)
    return paragraph


def validate_results(
    trials: pd.DataFrame, summary: pd.DataFrame, by_config: pd.DataFrame, include_live: bool
) -> None:
    if trials.duplicated(["task_id", "strategy"]).any():
        raise ValueError("Duplicate task-strategy trial pairs found.")
    expected_strategies = ["full_history", "paper_restatement", "json_state", "dependency_pruning"]
    expected_pairs = 60
    if include_live:
        expected_strategies.append("live_dependency_pruning")
        expected_pairs = 75
    if len(trials) != expected_pairs:
        raise ValueError(f"Expected exactly {expected_pairs} unique task-strategy trial pairs.")
    if trials["task_id"].nunique() != 15 or set(trials["strategy"].unique()) != set(expected_strategies):
        raise ValueError(f"Expected 15 fixtures and strategies {expected_strategies}.")
    counts = trials.groupby("strategy")["task_id"].nunique().to_dict()
    if counts != {strategy: 15 for strategy in expected_strategies}:
        raise ValueError(f"Each strategy must cover all 15 fixtures: {counts}")
    expected = {"full_history": 12, "paper_restatement": 13, "json_state": 15, "dependency_pruning": 15}
    base = trials[trials["strategy"].isin(expected)]
    actual = base.groupby("strategy")["success"].sum().astype(int).to_dict()
    if actual != expected:
        raise ValueError(f"Unexpected success totals: {actual}")
    connected = (
        by_config[
            (by_config["config_id"] == "connected_distractors")
            & (by_config["strategy"].isin(expected))
        ]
        .set_index("strategy")["successes"]
        .astype(int)
        .to_dict()
    )
    if connected != {"dependency_pruning": 5, "full_history": 3, "json_state": 5, "paper_restatement": 3}:
        raise ValueError(f"Unexpected connected-distractor results: {connected}")
    state = summary.set_index("strategy")["avg_state_bytes"]
    if round(float(state["json_state"]), 6) != round(float(state["dependency_pruning"]), 6):
        raise ValueError("JSON and dependency-pruning state sizes must match in the final panel.")


def load_report_data() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, Path, bool]:
    """Prefer the measured live-pruning supplement only when all 15 trials exist."""
    live_trial_path = LIVE_RESULTS / "processed" / "trial_results.csv"
    live_summary_path = LIVE_RESULTS / "processed" / "strategy_summary.csv"
    live_by_config_path = LIVE_RESULTS / "processed" / "strategy_by_config.csv"
    include_live = all(path.exists() for path in (live_trial_path, live_summary_path, live_by_config_path))
    results = LIVE_RESULTS if include_live else BASE_RESULTS
    trials = pd.read_csv(results / "processed" / "trial_results.csv")
    summary = pd.read_csv(results / "processed" / "strategy_summary.csv")
    by_config = pd.read_csv(results / "processed" / "strategy_by_config.csv")
    if include_live:
        live_trials = trials[trials["strategy"] == "live_dependency_pruning"]
        include_live = len(live_trials) == 15 and live_trials["task_id"].nunique() == 15
        if not include_live:
            results = BASE_RESULTS
            trials = pd.read_csv(results / "processed" / "trial_results.csv")
            summary = pd.read_csv(results / "processed" / "strategy_summary.csv")
            by_config = pd.read_csv(results / "processed" / "strategy_by_config.csv")
    validate_results(trials, summary, by_config, include_live)
    return trials, summary, by_config, results, include_live


def build() -> Path:
    trials, summary, by_config, results, include_live = load_report_data()
    metrics = summary.set_index("strategy")
    live_success = int(metrics.loc["live_dependency_pruning", "successes"]) if include_live else 0
    live_trials = int(metrics.loc["live_dependency_pruning", "trials"]) if include_live else 0
    live_state = float(metrics.loc["live_dependency_pruning", "avg_state_bytes"]) if include_live else 0.0
    json_state = float(metrics.loc["json_state", "avg_state_bytes"])
    live_reduction = 100 * (1 - live_state / json_state) if include_live and json_state else 0.0
    connected_live = 0
    if include_live:
        connected_live = int(
            by_config[
                (by_config["config_id"] == "connected_distractors")
                & (by_config["strategy"] == "live_dependency_pruning")
            ]["successes"].iloc[0]
        )
        live_tokens_per_success = float(metrics.loc["live_dependency_pruning", "tokens_per_success"])
        abstract_text = (
            "Reliable tool-using language models must infer valid call sequences while preserving intermediate "
            "values across multiple turns. StateBench reproduces a representative subset of Maekawa et al.'s "
            "FuncBenchGen, which generates synthetic hidden function-dependency graphs with controllable depth "
            "and distractors, and extends it by comparing five representations of observed working state. A "
            "paper-style Qwen3.6-27B calibration reproduced the qualitative depth effect but reached a "
            "zero-success floor on deep and connected tasks. I therefore applied one model-compatible interface "
            "consistently across 15 frozen graphs. The original four-strategy panel contained 60 executions: full "
            "history achieved 12/15 exact successes, paper restatement 13/15, and JSON plus static dependency "
            "pruning 15/15 each. Static pruning nevertheless retained the same 404.7 cumulative state bytes as "
            f"JSON. A post hoc, mechanism-driven live filter added 15 executions and achieved {live_success}/"
            f"{live_trials} while averaging {live_state:.1f} state bytes, a {live_reduction:.1f}% reduction from "
            "JSON. It remained perfect on easy and deep tasks but fell to 3/5 with connected distractors. The "
            "results therefore support explicit structured state while exposing a compression-reliability "
            "trade-off for aggressive public-schema pruning."
        )
        strategy_text = (
            "Full history adds no state suffix. Paper restatement appends every known value, and JSON serializes "
            "the same values as a compact object. Legacy dependency pruning retains values whose types are consumed "
            "anywhere in the visible schema set; every observed type remained globally consumable, producing the "
            "404.7-byte null result. Live dependency pruning instead computes a conservative backward frontier from "
            "the public target and type signatures, retains one non-conflicting value per still-needed exact type, "
            "does not retire failed calls, and emits no suffix after target completion. Neither pruning strategy "
            "consults hidden edges, node labels, ground-truth values, or the hidden target path."
        )
        result_text = (
            f"Across 75 primary executions, live dependency pruning achieved {live_success}/{live_trials}, compared with "
            "15/15 for JSON and legacy dependency pruning, 13/15 for paper restatement, and 12/15 for full history. "
            f"Its average cumulative state suffix was {live_state:.1f} bytes versus {json_state:.1f} for JSON "
            f"({live_reduction:.1f}% lower), and it used {live_tokens_per_success:,.0f} prompt tokens per success. "
            "Both live-pruning failures occurred on connected distractors and ended through reasoning-budget "
            "exhaustion. Prompt totals remain descriptive because failed trials can terminate early and hosted "
            "inference trajectories may differ even at temperature zero."
        )
        if live_state < json_state:
            pruning_discussion = (
                "The follow-up resolves the implementation cause of the original pruning null. The static rule "
                "asked whether a type was consumed anywhere, so stale values survived after their relevant consumer "
                "had completed and distractor-only consumers widened the allow-list. Live dependency pruning used "
                "only public schema information to retire types outside the still-unresolved target frontier. It "
                f"reduced average added-state bytes from {json_state:.1f} to {live_state:.1f} ({live_reduction:.1f}%). "
                f"That compression did not preserve reliability: it achieved {live_success}/{live_trials} overall "
                f"and {connected_live}/5 on connected distractors, versus 15/15 and 5/5 for JSON and static pruning. "
                "The result suggests that reducing external reminders can increase the model's burden to reconstruct "
                "the path internally. Because there are only five fixtures per condition, this trade-off is "
                "descriptive rather than statistically established."
            )
        else:
            pruning_discussion = (
                "The live-filter follow-up did not reduce average added-state bytes relative to JSON. Although it "
                "corrected the static global-consumer rule, the measured trajectories retained enough ambiguous "
                "public routes that no aggregate compression advantage emerged. This remains a negative result and "
                "does not support a pruning-efficiency claim."
            )
        reflection_text = (
            "This project taught me that reproducing an evaluation paper is not only about rerunning released code; "
            "the interaction contract between a benchmark and a model can determine whether the intended capability "
            "is measured at all. The most rewarding part was turning an initial all-failure local run into a "
            "controlled experiment without hiding negative findings or changing the core question. Rate limits, "
            "reasoning-budget exhaustion, and resumability were major implementation challenges. Frozen fixtures, "
            "viability gates, atomic checkpoints, and trace logging made the experiment recoverable. The static "
            "pruning null also taught me to inspect whether an intervention actually activates: rather than "
            "overclaiming, I traced the global-type rule, designed a public-schema liveness filter, and measured the "
            "follow-up separately. The activated filter demonstrated real compression, but its two connected-task "
            "failures reinforced that a smaller state is not automatically a better state."
        )
    else:
        abstract_text = (
            "Reliable tool-using language models must infer a valid sequence of calls while preserving intermediate values across multiple turns. Existing benchmarks often lack fine-grained difficulty controls and may be exposed to training or web-search contamination. StateBench reproduces a representative subset of Maekawa et al.'s FuncBenchGen, which generates synthetic hidden function-dependency graphs with controllable depth and distractors, and extends it by comparing four representations of already-observed working state. A paper-style Qwen3.6-27B calibration reproduced the qualitative depth effect but reached a zero-success floor on deep and connected tasks. I therefore applied one model-compatible interface consistently across all extension strategies and evaluated 15 frozen graphs under full history, paper restatement, compact JSON state, and dependency pruning (60 paired trials). Full history achieved 12/15 exact successes, paper restatement 13/15, and both compact strategies 15/15. The separation appeared under connected distractors, where compact state reached 5/5 versus 3/5 for both baselines. The results support explicit structured state, but the pruning filter retained the same state as JSON, so no independent pruning advantage was demonstrated."
        )
        strategy_text = (
            "Full history adds no state suffix. Paper restatement appends a readable list of all known values. JSON state serializes the same known values as a compact key-value object. Dependency pruning retains observed values whose visible types are consumed by at least one visible function schema; it never consults hidden edges or the target path. In this final panel every observed type remained consumable, so the pruning filter retained the same cumulative state as JSON. This is an experimental null result, not evidence of extra compression."
        )
        result_text = (
            "Across 60 paired trials, JSON and dependency pruning each achieved 15/15 successes, paper restatement 13/15, and full history 12/15. Full history has the lowest average prompt total, but its three failures ended early; fewer tokens therefore do not necessarily indicate a more efficient successful workflow. Tokens per success were 8,744 for dependency pruning, 8,887 for JSON, 9,394 for full history, and 9,958 for restatement."
        )
        pruning_discussion = (
            "The results also reject the narrow hypothesis that dependency pruning itself reduced context in this panel. JSON and pruning each averaged 404.7 cumulative state bytes because the schema-level filter removed nothing from these graphs. Their 143-token difference in average prompt totals follows different generated trajectories and call counts, not measured state compression. A stronger follow-up should include fixtures with observed types that become permanently unusable, repeat each model-strategy pair to estimate inference variability, increase the number of graph seeds, test more capable models under the paper-style interface, and report paired confidence intervals."
        )
        reflection_text = (
            "This project taught me that reproducing an evaluation paper is not only about rerunning released code; the interaction contract between a benchmark and a model can determine whether the intended capability is being measured at all. The most rewarding part was turning an initial all-failure local run into a controlled experiment without hiding the negative result or changing the central research question. Groq rate limits, reasoning-budget exhaustion, and provider-specific tool validation were the main implementation challenges. I addressed them with frozen fixtures, predeclared viability gates, atomic per-trial checkpoints, resumable execution, and trace-level error logging. Most importantly, I learned to treat null findings honestly: dependency pruning did not actually prune additional state in the final panel, but recognizing that limitation made the conclusion more precise and suggested a better next experiment."
        )

    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("StateBench: Working State for Reliable Multi-Step Tool Use")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Representative FuncBenchGen reproduction and controlled Qwen3.6-27B extension")
    metadata = doc.add_paragraph(style="Subtitle")
    metadata.paragraph_format.space_after = Pt(5)
    metadata.add_run("LebNet Tech Fellows 2026 | Academic Paper Option | [Your Name]")
    paragraph_border_bottom(metadata)

    doc.add_heading("1. Project Title & Abstract", level=1)
    add_abstract(
        doc,
        abstract_text,
    )

    doc.add_heading("2. Introduction & Problem Statement", level=1)
    body(
        doc,
        "Tool-augmented language models are increasingly expected to call APIs in workflows where later arguments depend on earlier outputs. A model can understand every function schema and still fail by selecting the wrong dependency path, calling a function before its inputs exist, or reusing an incorrect or stale value. Evaluation is difficult because realistic tool benchmarks are expensive to curate, expose only a few fixed difficulty levels, and can be contaminated when public tasks appear in training data or become searchable at test time.",
    )
    body(
        doc,
        "FuncBenchGen addresses these limitations by casting tool use as traversal over a hidden directed acyclic graph (DAG). Nodes are synthetic single-output functions; edges encode type-compatible dependencies. The model sees initial values, randomized tool schemas, and a target variable, but not the graph. Difficulty can be controlled through required core functions, maximum dependency depth, and irrelevant functions. Connected irrelevant nodes (CINs) are especially challenging because they share compatible types with the solution graph. The paper also reports that value-not-yet-known errors dominate execution traces and proposes restating all known variables after each call. StateBench asks the same core question while adding a focused extension: when tasks, model, schemas, call budget, and seeds are fixed, does the representation of observed working state change reliability and context cost?",
    )
    add_comparison_table(doc)

    page_break(doc)
    doc.add_heading("3. Methodology (Approach / Reproduction Details)", level=1)
    body(
        doc,
        "Reproduction foundation. The project vendors the official FuncBenchGen code release at commit 0718d1cf25b601d0b25fbbbbd064525536cea876. It preserves graph generation, hidden dependencies, type/subtype matching, deterministic target values, exact-answer evaluation, five graph seeds per setting, randomized function order, and the paper's maximum call budget of twice the minimum required calls. No external static dataset is used: graphs are generated synthetically and then frozen so every state strategy receives the identical graph, functions, initial values, target, and tool ordering.",
    )
    body(
        doc,
        "The paper-style calibration used opaque descriptions and the upstream behavior in which wrong inputs execute and return an incorrect value. With full history, Qwen3.6-27B achieved 2/5 on easy depth, 0/5 on deep depth, and 0/5 with connected distractors. This reproduces the qualitative depth decline, but the zero deep baseline prevents a faithful estimate of CINs' additional effect. It is therefore reported separately from the final extension rather than mixed with it.",
    )
    body(
        doc,
        "For the controlled extension, every strategy used the same model-compatible interface: descriptions exposed variable names, premature calls returned an explicit error, the model made one tool call per response, and it was instructed to use only observed values. The hidden DAG, target path, ground-truth values, correctness labels, and unobserved state remained unavailable. Qwen3.6-27B was accessed through Groq with temperature 0, top-p 1, default reasoning enabled, a 4,096-token completion allowance, and the same two-times-minimum call budget.",
    )
    add_condition_table(doc)
    labeled_body(doc, "State strategies.", strategy_text)
    labeled_body(doc, "Measures.", "Primary outcomes are exact target success and total request-level prompt tokens. Secondary measures are completion tokens, tool calls, prompt tokens per success, success per 1,000 prompt tokens, cumulative bytes added by state suffixes, and trace-level failure codes. A trial may contain a recoverable intermediate error and still succeed, so error occurrences are not mutually exclusive final outcomes.")

    page_break(doc)
    doc.add_heading("4. Implementation Details & Results", level=1)
    body(
        doc,
        "StateBench is implemented in Python 3.11 around the pinned generator/evaluator, with NetworkX for graph reconstruction, an OpenAI-compatible Groq provider, pandas for aggregation, and Matplotlib for figures. Each raw trial is written atomically as JSON. Resume logic skips completed pairs, short-window rate limits are retried without discarding the active conversation, and daily quota exhaustion stops safely for a later run. Stored traces include requests, tool calls, arguments, results, token counts, state suffix size, and canonical failure codes.",
    )
    add_results_table(doc, summary, include_live)
    figure1 = results / "figures" / "reliability_by_configuration.png"
    add_picture(doc, figure1, 6.45, "Grouped bar chart of exact success by task configuration and state strategy")
    if include_live:
        add_caption(doc, f"Figure 1. Exact success by condition across five strategies; live pruning reached {connected_live}/5 on connected distractors.")
        figure2 = results / "figures" / "live_pruning_comparison.png"
        add_picture(doc, figure2, 6.65, "Two-panel comparison of compact-state reliability and cumulative state bytes")
        add_caption(doc, "Figure 2. Reliability and added-state size for JSON, static dependency pruning, and live dependency pruning. Bytes measure cumulative suffix content added during a trial.")
    else:
        add_caption(doc, "Figure 1. Exact success by condition. Connected distractors create the clearest separation: compact state 5/5 versus 3/5 for full history and restatement.")
        figure2 = results / "figures" / "context_and_trace_errors.png"
        add_picture(doc, figure2, 6.65, "Two-panel chart of reliability versus prompt tokens and trace-level error occurrences")
        add_caption(doc, "Figure 2. Left: aggregate reliability-context trade-off. Right: trace errors (teal: value not yet known; red: thinking budget exhausted). Recoverable events are not counts of final failed trials.")
    body(
        doc,
        result_text,
        after=2,
    )

    doc.add_heading("5. Discussion & Analysis", level=1)
    body(
        doc,
        "The strongest result is conditional rather than universal. On easy tasks, full history failed one seed while every original explicit-state strategy succeeded. On deep tasks without distractors, the original four strategies reached 5/5. Adding ten connected irrelevant functions reduced full history to 3/5, reproducing the paper's qualitative claim that type-compatible distractors confuse dependency selection. JSON and static dependency pruning remained 5/5, recovering both connected full-history failures and the single easy failure."
        + (f" Live dependency pruning achieved {live_success}/15 overall and {connected_live}/5 on connected tasks." if include_live else ""),
    )
    body(
        doc,
        "Paper restatement produced only a small aggregate gain: 13/15 versus 12/15. In the connected condition it succeeded on a full-history failure at seed 300 but failed a different seed 100, while both failed seed 400; its net result remained 3/5. This qualifies the paper's mitigation finding. Simply repeating every value was not sufficient for this model under the stress condition, whereas a compact machine-readable map was. Because the sample is small, this should be treated as evidence for further testing rather than proof that JSON formatting is intrinsically superior.",
    )
    body(
        doc,
        "Trace behavior clarifies the mechanism. Sixteen value-not-yet-known events and seven thinking-budget exhaustion events occurred across the five strategies, echoing the paper's state-availability and reasoning-capacity concerns. Under the compatible interface, the executor exposed premature calls instead of silently propagating a fabricated value, and compact-state runs sometimes recovered. The unsuccessful live-pruning runs both reached the 4,096-token per-turn allowance without producing another tool call or the target. This is a valid outcome under the fixed protocol, but it does not prove that pruning alone caused the looping behavior. Absolute success rates cannot be compared directly with the paper's model table.",
    )
    body(
        doc,
        pruning_discussion,
    )

    doc.add_heading("6. Reflection on Learnings", level=1)
    body(
        doc,
        reflection_text,
        after=4,
    )

    doc.add_heading("References", level=1)
    body(
        doc,
        "[1] Maekawa, S., Hassell, J., Pezeshkpour, P., Mitchell, T., and Hruschka, E. (2026). Towards Reliable Benchmarking: A Contamination Free, Controllable Evaluation Framework for Multi-step LLM Function Calling. ICLR 2026; arXiv:2509.26553v2.",
        after=1.5,
    )
    body(
        doc,
        "[2] Megagon Labs. FuncBenchGen code release, commit 0718d1cf25b601d0b25fbbbbd064525536cea876. Model: qwen/qwen3.6-27b via Groq.",
        after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
